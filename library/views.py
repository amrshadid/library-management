import multiprocessing
from django.template.loader import get_template, render_to_string
from io import BytesIO
import cgi
from io import StringIO
# from reportlab.pdfgen import canvas
from django.template import Context
from django.shortcuts import render
import os
from xhtml2pdf import pisa
from django.http import HttpResponse
from django.conf import settings
from django.core.exceptions import ObjectDoesNotExist
from django.http import Http404
from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from rest_framework.generics import (
    ListAPIView, RetrieveAPIView, CreateAPIView,
    UpdateAPIView, DestroyAPIView
)
from rest_framework.authentication import TokenAuthentication
from rest_framework.permissions import AllowAny, IsAuthenticated, IsAuthenticatedOrReadOnly, IsAdminUser
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.status import (HTTP_200_OK,
                                   HTTP_400_BAD_REQUEST,
                                   HTTP_201_CREATED,
                                   HTTP_204_NO_CONTENT
                                   )
from .serializers import (
    IssueSerializer,
    IssueDetailSerializer,
    RuleSerializer,
    UserSelectedIssueSerializer,
    SubCategorySerializer
)
from .models import *
from users.models import CustomUser, Teams
from checkout.models import StripeCustomer
from rest_framework import viewsets
from django.db.models import Q
from django.views.decorators.csrf import csrf_exempt
from rest_framework.authtoken.models import Token
from django.views.generic import View
# Create your views here.
from master_library.settings import BASE_DIR, SERVER_CREATED_FILES_PATH
from rest_framework.authentication import BasicAuthentication
from PyPDF2 import PdfFileReader
from uuid import uuid4
from docx import Document

URL_WEB='library-mp1.herokuapp'

class homeView(View):
    def get(self, request):
        try:
            with open(os.path.join(BASE_DIR, 'frontend', 'out', 'index.html')) as file:
                return HttpResponse(file.read())
        except:
            return HttpResponse(
                """
                index.html not found ! build your React app !!
                """,
                status=501,
            )


class getSelectedRule(APIView):
    def get(self, request):
        user_id = Token.objects.get(
            key=request.GET[URL_WEB]).user_id
        user_instance = CustomUser.objects.get(id=user_id)
        qs = UserSelectedIssue.objects.filter(user=user_instance)
        data = {}
        mapping_data = {}
        mapping_qs = Mapping.objects.filter()
        for mq in mapping_qs:
            if mq.issue.category in mapping_data:
                if mq.subCategory.subCategory in mapping_data[mq.issue.category]:
                    mapping_data[mq.issue.category][mq.subCategory.subCategory].append(
                        mq.issue.title)
                else:
                    mapping_data[mq.issue.category][mq.subCategory.subCategory] = [
                        mq.issue.title]
            else:
                mapping_data[mq.issue.category] = {
                    mq.subCategory.subCategory: [
                        mq.issue.title]
                }

        for k in qs:
            if k.rule.issue.title in data.keys():
                data[k.rule.issue.title].append(k.rule.rule)
            else:
                data[k.rule.issue.title] = [k.rule.rule]
        #qs = qs.values('rule').annotate()
        # print(qs)
        # print(data)
        serializer = UserSelectedIssueSerializer(qs, many=True)
        return Response({'status': 1, 'message': 'Success', 'data': data, "mappingData": mapping_data})


class SubCategoryView(APIView):
    permission_classes = (IsAuthenticated,)
    # authentication_classes = (TokenAuthentication,)
    authentication_classes = [BasicAuthentication]

    def get(self, request):
        queryset = SubCategory.objects.all()
        serializer = SubCategorySerializer(queryset, many=True)
        return Response({'status': 1, 'message': 'Success', 'subCategories': serializer.data})

    @csrf_exempt
    def post(self, request):
        serializer = SubCategorySerializer(data=request.data)
        if(serializer.is_valid()):
            serializer.save()
            return Response(serializer.data, status=HTTP_201_CREATED)
        return Response(serializer.errors, status=HTTP_400_BAD_REQUEST)


class CustomAdminView(APIView):
    #permission_classes = (IsAdminUser,)
    authentication_classes = [BasicAuthentication]

    @csrf_exempt
    def post(self, request, category):
        # retrieving list of rules
        user_instance = Token.objects.get(
            key=request.POST[URL_WEB]).user
        rule_title = request.data.getlist('rules')
        rule_isState = request.data.getlist('isStateSpecific')
        rule_state = request.data.getlist('state')
        sub_category = request.data.getlist('subCategories')
        # for creating issue object
        serializer = IssueSerializer(data=request.data)
        if(serializer.is_valid()):
            serializer.save(author=user_instance, category=category)
            issue_instance = serializer.save()
            n = len(rule_title)
            for i in range(n):
                isSpecific = True
                if((rule_state[i]) == ""):
                    isSpecific = False
                rule_obj = Rule(rule=rule_title[i], isStateSpecific=isSpecific, state=rule_state[i],
                                author=user_instance, issue=issue_instance)
                rule_obj.save()
            n = len(sub_category)
            for k in range(n):
                cat_obj = SubCategory.objects.get(id=int(sub_category[k]))
                print(cat_obj)
                obj = Mapping(issue=issue_instance, subCategory=cat_obj)
                obj.save()
            # to return detailed view of issue created
            output_data = IssueDetailSerializer(issue_instance)
            return Response({'status': 1, 'message': 'Added Successfully', 'data': output_data.data})
        return Response({'status': 0, 'message': 'Request Failed', 'data': serializer.errors})


class AdminIssuesListView(APIView):
    def get(self, request):
        user_instance = None
        queryset = Issue.objects.filter(iscustom=False)
        serializer = IssueSerializer(queryset, many=True)
        return Response({'status': 1, 'message': 'Success', 'issueList': serializer.data})


class CustomIssueEditView(APIView):
    #permission_classes = (IsAdminUser,)
    authentication_classes = [BasicAuthentication]

    @csrf_exempt
    def post(self, request):
        # retrieving list of rules
        user_instance = Token.objects.get(
            key=request.POST[URL_WEB]).user
        rule_title = request.data.getlist('rules')
        rule_isState = request.data.getlist('isStateSpecific')
        rule_state = request.data.getlist('state')
        issueData = Issue.objects.get(id=int(request.POST['title']))
        n = len(rule_title)
        for i in range(n):
            isSpecific = True
            if((rule_state[i]) == ""):
                isSpecific = False
            rule_obj = Rule(rule=rule_title[i], isStateSpecific=isSpecific, state=rule_state[i],
                            author=user_instance, issue=issueData)
            rule_obj.save()
        output_data = IssueDetailSerializer(issueData)
        return Response({'status': 1, 'message': 'Added Successfully', 'data': output_data.data})


class IssueListView(APIView):
    #permission_classes = (IsAuthenticated,)
    #authentication_classes = (TokenAuthentication,)
    authentication_classes = [BasicAuthentication]

    def get(self, request):
        category = request.GET['category']
        sub_category = request.GET['subCategory']
        try:
            user_id = Token.objects.get(
                key=request.GET[URL_WEB]).user_id
            user_instance = CustomUser.objects.get(id=user_id)
            # MISSING NECESSARY CHECKS FOR PAYMENT PLAN REL. ISSUES RENDERING
            user_payment_info = StripeCustomer.objects.filter(
                user__id=user_instance.id)
            if user_payment_info:
                isTrailPlan = True if user_payment_info[0].plan == "TR" else False
            else:
                isTrailPlan = False
        except Exception as e:
            print(e)
            print("in heree")
            user_instance = None
        print(user_instance)
        if sub_category.isdigit() and int(sub_category) != 0:
            sub_category_instance = SubCategory.objects.get(
                id=int(sub_category))
            mapping_qs = Mapping.objects.filter(
                subCategory=sub_category_instance)
            if(request.GET['plan'] == "TR"):
                mapping_qs = mapping_qs.exclude(issue__is_free=False)
            if category == "AR":
                mapping_qs = mapping_qs.exclude(issue__category="RR")
            else:
                mapping_qs = mapping_qs.exclude(issue__category="AR")

            if user_instance and user_instance.is_staff is False:
                print("yyyyyy")
                mapping_qs = (mapping_qs.filter(issue__iscustom=True, issue__author=user_instance) |
                              mapping_qs.filter(issue__iscustom=False))
            else:
                mapping_qs = mapping_qs.filter(issue__iscustom=False)

            queryset = Issue.objects.none()
            for qs in mapping_qs:
                instance = qs.issue
                queryset |= Issue.objects.filter(pk=instance.pk)
        else:
            if category == "AR":
                queryset = Issue.objects.exclude(category="RR")
            else:
                queryset = Issue.objects.exclude(category="AR")

            if user_instance and user_instance.is_staff == False:
                queryset = (queryset.filter(iscustom=True, author=user_instance) |
                            queryset.filter(iscustom=False))
            else:
                queryset = queryset.filter(iscustom=False)
            if isTrailPlan:
                queryset = queryset.filter(is_free=True)
        serializer = IssueSerializer(queryset, many=True)
        return Response({'status': 1, 'message': 'Success', 'issueList': serializer.data})

    @csrf_exempt
    def post(self, request):
        user_id = Token.objects.get(
            key=request.POST[URL_WEB]).user_id
        user_instance = CustomUser.objects.get(id=user_id)
        # ----------- check for trail plan and issues addition -------
        user_payment_info = StripeCustomer.objects.get(
            user__id=user_instance.id)
        isTrailPlan = True if user_payment_info.plan == "TR" else False
        if isTrailPlan:
            category = request.POST['category']
            if category == "AR":
                queryset = Issue.objects.exclude(category="RR")
            else:
                queryset = Issue.objects.exclude(category="AR")
            # print(list(val.category for val in queryset))
            queryset = queryset.filter(
                iscustom=True, author=user_instance, is_free=True)
            # print(len(queryset))
            if len(queryset) >= 4:
                return Response({'status': 0, 'message': 'You can create maximum of 4 issues with trail plan', 'data': None})

        # ----------- check for trail plan and issues addition -------

        rule_title = request.data.getlist('rules')
        # rule_isState = request.data.getlist('isStateSpecific')
        rule_state = request.data.getlist('state')
        if('subCategories' in request.POST):
            sub_category = request.POST['subCategories']
        category = request.POST['category']
        serializer = IssueSerializer(data=request.data)
        if(serializer.is_valid()):
            #serializer.save(author = self.request.user, category = category, iscustom = True)
            serializer.save(author=user_instance,
                            category=category, iscustom=True, is_free=isTrailPlan)
            issue_instance = serializer.save()
            n = len(rule_title)
            for i in range(n):
                isSpecific = True
                if((rule_state[i]) == ""):
                    isSpecific = False
                rule_obj = Rule(rule=rule_title[i], isStateSpecific=isSpecific, state=rule_state[i],
                                iscustom=True, author=user_instance, issue=issue_instance)
                rule_obj.save()
            if('subCategories' in request.POST):
                cat_obj = SubCategory.objects.get(id=int(sub_category))
                obj = Mapping(issue=issue_instance, subCategory=cat_obj)
                obj.save()
            # to return detailed view of issue created
            output_data = IssueDetailSerializer(issue_instance)
            return Response({'status': 1, 'message': 'Added Successfully', 'data': output_data.data})
        return Response({'status': 0, 'message': 'Request Failed', 'data': serializer.errors})


class IssueDetailView(RetrieveAPIView):
    authentication_classes = [BasicAuthentication]

    def get_object(self, pk):
        # Returns an object instance that should
        # be used for detail views.
        try:
            return Issue.objects.get(pk=pk)
        except Issue.DoesNotExist:
            raise Http404

    def get(self, request, pk):
        pk = int(pk)
        queryset = Issue.objects.get(id=pk)
        serializer = IssueDetailSerializer(queryset)
        #request.session['selected_issue_id'] = pk
        return Response({'status': 1, 'message': 'Success', 'issue': serializer.data})

    def put(self, request, pk, format=None):
        try:
            user_id = Token.objects.get(key=request.POST['userId']).user_id
            user_instance = CustomUser.objects.get(id=user_id)
        except:
            return HTTP_400_BAD_REQUEST
        issue_obj = self.get_object(pk)
        if issue_obj.author == user_instance:
            serializer = IssueSerializer(issue_obj, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=HTTP_400_BAD_REQUEST)
        else:
            return HTTP_400_BAD_REQUEST

    def patch(self, request, pk, format=None):
        try:
            user_id = Token.objects.get(key=request.POST['userId']).user_id
            user_instance = CustomUser.objects.get(id=user_id)
        except:
            return HTTP_400_BAD_REQUEST
        issue_obj = self.get_object(pk)
        if issue_obj.author == user_instance:
            serializer = IssueSerializer(
                issue_obj, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=HTTP_400_BAD_REQUEST)
        else:
            return HTTP_400_BAD_REQUEST

    def delete(self, request, pk, format=None):
        try:
            user_id = Token.objects.get(key=request.POST['userId']).user_id
            user_instance = CustomUser.objects.get(id=user_id)
        except:
            return HTTP_400_BAD_REQUEST
        issue_obj = self.get_object(pk)
        if issue_obj.author == user_instance:
            issue_obj.delete()
            return Response(status=HTTP_204_NO_CONTENT)


class RuleView(APIView):
    authentication_classes = [BasicAuthentication]

    def check_new(self, user_obj, rule_obj):
        if not user_obj:
            return 0
        if (rule_obj.created - user_obj.last_login).days > 1:
            return 1
        else:
            return 0

    def check_selected(self, user_obj, rule_obj):
        if not user_obj:
            return 0
        selected_issue = UserSelectedIssue.objects.filter(
            rule=rule_obj, user=user_obj)
        if selected_issue:
            return 1
        else:
            return 0

    def get(self, request):
        try:
            user_id = Token.objects.get(
                key=request.GET[URL_WEB]).user_id
            user_instance = CustomUser.objects.get(id=user_id)
        except:
            user_instance = None
        print(user_instance)
        selected_issue_id = int(request.GET['issue_id'])
        issueData = Issue.objects.get(id=selected_issue_id)
        queryset = Rule.objects.filter(
            issue__id=selected_issue_id, iscustom=False)
        if user_instance and user_instance.is_staff == False:
            queryset |= Rule.objects.filter(
                issue__id=selected_issue_id, author=user_instance)

        data = []
        for v in queryset:
            conflictingRulesData = RuleMapping.objects.filter(rule=v)
            conflictData = []
            for k in conflictingRulesData:
                conflictTemp = {
                    'rule': k.rule_mapped.rule
                }
                conflictData.append(conflictTemp)
            temp = {
                'id': v.id,
                'rule': v.rule,
                'isStateSpecific': v.isStateSpecific,
                'state': v.state,
                'is_new': self.check_new(user_instance, v),
                'is_selected': self.check_selected(user_instance, v),
                'conflicting_rules': conflictData,
            }
            data.append(temp)
        return Response({'status': 1, 'message': 'Success', 'ruleList': data, 'issueNote': issueData.issue_note})

    def post(self, request):
        #selected_issue_id = int(request.session.get('selected_issue_id'))
        selected_issue_id = int(request.POST['issue_id'])
        user_id = Token.objects.get(
            key=request.POST[URL_WEB]).user_id
        user_instance = CustomUser.objects.get(id=user_id)
        serializer = RuleSerializer(data=request.data)
        issue_instance = Issue.objects.filter(id=selected_issue_id).first()
        if(serializer.is_valid()):
            if user_instance.is_staff == False:
                serializer.save(
                    iscustom=True, issue=issue_instance, author=user_instance)
                return Response({'status': 1, 'message': 'Rule Added Successfully'})
            else:
                serializer.save(issue=issue_instance, author=user_instance)
            return Response({"message": "Rule Created"}, serializer.data, status=HTTP_201_CREATED)
        return Response(serializer.errors, status=HTTP_400_BAD_REQUEST)


class RuleDetailView(APIView):
    authentication_classes = [BasicAuthentication]

    def get_object(self, pk):
        # Returns an object instance that should
        # be used for detail views.
        try:
            return Rule.objects.get(pk=pk)
        except Rule.DoesNotExist:
            raise Http404

    def put(self, request, pk, format=None):
        try:
            user_id = Token.objects.get(key=request.POST['userId']).user_id
            user_instance = CustomUser.objects.get(id=user_id)
        except:
            return HTTP_400_BAD_REQUEST
        rule_obj = self.get_object(pk)
        if rule_obj.author == user_instance:
            serializer = IssueSerializer(rule_obj, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=HTTP_400_BAD_REQUEST)
        else:
            return HTTP_400_BAD_REQUEST

    def patch(self, request, pk, format=None):
        try:
            user_id = Token.objects.get(key=request.POST['userId']).user_id
            user_instance = CustomUser.objects.get(id=user_id)
        except:
            return HTTP_400_BAD_REQUEST
        rule_obj = self.get_object(pk)
        if rule_obj.author == user_instance:
            serializer = IssueSerializer(
                rule_obj, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=HTTP_400_BAD_REQUEST)
        else:
            return HTTP_400_BAD_REQUEST

    def delete(self, request, pk, format=None):
        try:
            user_id = Token.objects.get(key=request.POST['userId']).user_id
            user_instance = CustomUser.objects.get(id=user_id)
        except:
            return HTTP_400_BAD_REQUEST
        rule_obj = self.get_object(pk)
        if rule_obj.author == user_instance:
            rule_obj.delete()
            return Response(status=HTTP_204_NO_CONTENT)
        else:
            return HTTP_400_BAD_REQUEST


class UserSelectedIssueView(APIView):
    authentication_classes = [BasicAuthentication]

    def get(self, request):
        user_id = Token.objects.get(
            key=request.GET[URL_WEB]).user_id
        user_instance = CustomUser.objects.get(id=user_id)
        qs = UserSelectedIssue.objects.filter(user=user_instance)
        #qs = qs.values('rule').annotate()
        # print(qs)
        serializer = UserSelectedIssueSerializer(qs, many=True)
        # print(serializer)
        return Response({'status': 1, 'message': 'Success', 'selected_issue': serializer.data})

    def post(self, request):
        ruleId = request.POST['ruleId']
        ruleId = int(ruleId.split('-')[1])
        print(ruleId)
        user_instance = Token.objects.get(
            key=request.POST[URL_WEB]).user
        rule_instance = Rule.objects.get(id=ruleId)
        if rule_instance is None:
            return Response({"message": "Invalid Data"}, status=HTTP_400_BAD_REQUEST)
        checkData = UserSelectedIssue.objects.filter(
            user=user_instance, rule=rule_instance)
        if(checkData):
            checkData = UserSelectedIssue.objects.get(
                user=user_instance, rule=rule_instance)
            checkData.delete()
        else:
            obj = UserSelectedIssue(user=user_instance, rule=rule_instance)
            obj.save()
        return Response({'status': 1, "message": "Rules Saved Successfully"}, status=HTTP_201_CREATED)

    def delete(self, request):
        rule_id = request.data.get('rule')
        rule_instance = Rule.objects.get(id=rule_id)
        if rule_instance is None:
            return Response({"message": "Invalid Data"}, status=HTTP_400_BAD_REQUEST)
        obj = UserSelectedIssue.objects.get(
            user=self.request.user, rule=rule_instance)
        if obj:
            obj.delete()
            return Response(status=HTTP_204_NO_CONTENT)
        else:
            return HTTP_400_BAD_REQUEST


class getSubCategories(APIView):
    def get(self, request, category_name):
        if(category_name == "AR"):
            subCategoryData = SubCategory.objects.filter(
                subCategory__startswith="Architectural")
        elif(category_name == "RR"):
            subCategoryData = SubCategory.objects.exclude(
                subCategory__startswith="Architectural")
        else:
            subCategoryData = SubCategory.objects.all()
        serializer = SubCategorySerializer(subCategoryData, many=True)
        return Response({'status': 1, 'data': serializer.data})
# class SubAouRulesView(RetrieveAPIView):

#     def get_object(self,subAou):
#         try:
#             subAou_data = UserSelectedIssue.objects.filter(user=subAou)
#             return subAou_data
#         except ObjectDoesNotExist:
#             subAou_data = UserSelectedIssue.objects.none()
#             return subAou_data

#     def get(self,request):
#         user_id = Token.objects.get(key= request.GET[URL_WEB]).user_id
#         user_instance = CustomUser.objects.get(id=user_id)
#         if user_instance.is_aou == True:
#             subAou_list = Teams.objects.filter(owners = user_instance)
#             if subAou_list:
#                 data = []
#                 for subAou in subAou_list:
#                     selected_rules = self.get_object(subAou.subAou)
#                     rules = RuleSerializer(selected_rules,many=True)
#                     temp = {
#                         'user': subAou.subAou.email,
#                         'rules': rules.data
#                     }
#                     data.append(temp)
#                 return Response({'message':'Success','selected_issue':data}, status=HTTP_200_OK)
#             return Response({"message":"No members in Team"}, status=HTTP_400_BAD_REQUEST)
#         return Response({"message":"Not Authorised"}, status=HTTP_400_BAD_REQUEST)


class SubAouRulesView(APIView):
    authentication_classes = [BasicAuthentication]

    def get(self, request):
        user_id = Token.objects.get(
            key=request.GET[URL_WEB]).user_id
        user_instance = CustomUser.objects.get(id=user_id)
        if user_instance.is_aou == True:
            subAou_list = Teams.objects.filter(owners=user_instance)
        else:
            return Response({'status': 0, 'message': 'Login with Aou account'})
        issueList = Issue.objects.all()
        data = []
        for k in issueList:
            ruleList = Rule.objects.filter(issue=k)
            ruleData = []
            for x in ruleList:
                subAouRuleStatus = []
                for m in subAou_list:
                    ruleStatus = UserSelectedIssue.objects.filter(
                        user=m.subAou, rule=x)
                    if(ruleStatus):
                        subAouRuleStatus.append(
                            {'subAouId': m.id, 'ruleStatus': True})
                    else:
                        subAouRuleStatus.append(
                            {'subAouId': m.id, 'ruleStatus': False})
                ruleStatus = UserSelectedIssue.objects.filter(
                    user=user_instance, rule=x)
                if(ruleStatus):
                    temp = {
                        'ruleId': x.id,
                        'rule': x.rule,
                        'subAouRuleStatus': subAouRuleStatus,
                        'Aou': True
                    }
                    # subAouRuleStatus.append({'AouId':user_instance.id,'ruleStatus':True})
                else:
                    temp = {
                        'ruleId': x.id,
                        'rule': x.rule,
                        'subAouRuleStatus': subAouRuleStatus,
                        'Aou': False
                    }
                    # subAouRuleStatus.append({'AouId':user_instance.id,'ruleStatus':False})
                # temp={
                #     'ruleId':x.id,
                #     'rule':x.rule,
                #     'subAouRuleStatus':subAouRuleStatus
                # }
                ruleData.append(temp)
            data.append(
                {'issueId': k.id, 'issue': k.title, 'ruleData': ruleData})
        return Response({'status': 1, 'data': data})


def fetch_resources(uri, rel):
    path = os.path.join(uri.replace(settings.STATIC_URL, ""))
    return path


def combine_word_documents(files):
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
        # if index < len(files)-1:
        #     sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
    filepath = os.path.join(SERVER_CREATED_FILES_PATH, str(uuid4())+".docx")
    merged_document.save(filepath)
    return filepath


def convert_pdf_to_docx(pdf_path, s_i, e_i):
    filename_docx = os.path.join(
        SERVER_CREATED_FILES_PATH, str(uuid4())+".docx")
    parse(pdf_path, filename_docx, start=s_i, end=e_i, multiprocessing=True)
    return filename_docx


def render_to_pdf(template_src, isDocx, data):

    template = get_template(template_src)
    #context = Context(data)
    html = template.render(data)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("utf-8")), result)
    if isDocx:
        available_docxs = []
        filename_pdf = os.path.join(
            SERVER_CREATED_FILES_PATH, str(uuid4())+".pdf")
        with open(filename_pdf, "wb") as f:
            f.write(result.getbuffer())
        pdf_pages = PdfFileReader(result).getNumPages()
        for i in range(pdf_pages):
            docx_path = convert_pdf_to_docx(filename_pdf, i, i+1)
            available_docxs.append(docx_path)
        final_merged_docx = combine_word_documents(available_docxs)
        with open(final_merged_docx, "rb") as f:
            result = BytesIO(f.read())
        for docx_rm in available_docxs:
            os.remove(docx_rm)
        os.remove(final_merged_docx)
        os.remove(filename_pdf)
    return result.getvalue()


class GeneratePDF(APIView):
    def get(self, request, *args, **kwargs):
        try:
            key = request.GET.get(URL_WEB)
            # print(key)
            user_id = Token.objects.get(key=key).user_id
            user_instance = CustomUser.objects.get(id=user_id)
            selected_rule = UserSelectedIssue.objects.filter(
                user=user_instance)  # you can filter using order_id as well
        except:
            return HttpResponse("505 Not Found")
        data_ar = {}
        data_rr = {}
        for sr in selected_rule:
            rule = sr.rule
            issue = rule.issue

            if issue.category == 'AR' or issue.category == 'ARR':
                if issue.title in data_ar:
                    data_ar[issue.title].append(rule.rule)
                else:
                    data_ar[issue.title] = [rule.rule]
            if issue.category == 'RR' or issue.category == 'ARR':
                if issue.title in data_rr:
                    data_rr[issue.title].append(rule.rule)
                else:
                    data_rr[issue.title] = [rule.rule]

        op_ar = ''
        for k, v in data_ar.items():
            temp = ''
            temp += '<h2 style="font-size:20px">'+k + \
                '</h2> <div style="margin-left :  20px ;"> <ul style="list-style-type:none;">'
            for j in v:
                temp += '<li style="font-size:15px"> => &nbsp;'+j+'</li>'
            temp += '</ul> </div>'
            op_ar += temp
        op_rr = ''
        for k, v in data_rr.items():
            temp = ''
            temp += '<h2 style="font-size:20px">'+k + \
                '</h2> <div style="margin-left :  20px ;"> <ul style="list-style-type:none;">'
            for j in v:
                temp += '<li style="font-size:15px"> => &nbsp;'+j+'</li>'
            temp += '</ul> </div>'
            op_rr += temp

        pdf = render_to_pdf('issue_doc.html', False, {
                            'association': user_instance.association_name, 'arch_gdl': op_ar, 'reg_rl': op_rr})

        return HttpResponse(pdf, content_type='application/pdf')


class GenerateDocx(APIView):
    def get(self, request, *args, **kwargs):
        try:
            key = request.GET.get(URL_WEB)
            # print(key)
            user_id = Token.objects.get(key=key).user_id
            user_instance = CustomUser.objects.get(id=user_id)
            selected_rule = UserSelectedIssue.objects.filter(
                user=user_instance)  # you can filter using order_id as well
        except:
            return HttpResponse("505 Not Found")
        data_ar = {}
        data_rr = {}
        for sr in selected_rule:
            rule = sr.rule
            issue = rule.issue

            if issue.category == 'AR' or issue.category == 'ARR':
                if issue.title in data_ar:
                    data_ar[issue.title].append(rule.rule)
                else:
                    data_ar[issue.title] = [rule.rule]
            if issue.category == 'RR' or issue.category == 'ARR':
                if issue.title in data_rr:
                    data_rr[issue.title].append(rule.rule)
                else:
                    data_rr[issue.title] = [rule.rule]

        op_ar = ''
        for k, v in data_ar.items():
            temp = ''
            temp += '<h2 style="font-size:20px">'+k + \
                '</h2> <div style="margin-left :  20px ;"> <ul style="list-style-type:none;">'
            for j in v:
                temp += '<li style="font-size:15px"> => &nbsp;'+j+'</li>'
            temp += '</ul> </div>'
            op_ar += temp
        op_rr = ''
        for k, v in data_rr.items():
            temp = ''
            temp += '<h2 style="font-size:20px">'+k + \
                '</h2> <div style="margin-left :  20px ;"> <ul style="list-style-type:none;">'
            for j in v:
                temp += '<li style="font-size:15px"> => &nbsp;'+j+'</li>'
            temp += '</ul> </div>'
            op_rr += temp

        pdf = render_to_pdf('issue_doc.html', True, {
                            'association': user_instance.association_name, 'arch_gdl': op_ar, 'reg_rl': op_rr})

        response = HttpResponse(
            pdf, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="rules.docx"'
        return response
